import os
import sys
import unittest

SCRIPTS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "scripts"
)
sys.path.insert(0, SCRIPTS_DIR)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document  # noqa: E402
from render_coverletter_docx import render_coverletter_docx  # noqa: E402


def _minimal_letter_data(**overrides):
    data = {
        "company_name": "Acme Corp",
        "tagline": "PRODUCT MANAGER | GROWTH",
        "greeting": "Dear Acme Corp Hiring Team,",
        "body_paragraphs": [
            "I'm excited to apply for this role at Acme Corp.",
            "My background lines up well with what you need.",
        ],
        "sign_off": "Sincerely,",
    }
    data.update(overrides)
    return data


class TestRenderCoverletterDocx(unittest.TestCase):

    def setUp(self):
        # See test_render_coverletter.py: sandboxed so these assert on the
        # renderer, not on the operator's own contact details.
        import persona

        self._sandbox = persona.sandbox_profile()
        self._sandbox.__enter__()
        self.out_path = os.path.join(
            os.path.dirname(__file__), "_tmp_coverletter_docx_test.docx"
        )

    def tearDown(self):
        self._sandbox.__exit__(None, None, None)
        if os.path.exists(self.out_path):
            os.remove(self.out_path)

    def _paragraph_texts(self, doc):
        return [p.text for p in doc.paragraphs]

    def test_header_contains_senders_name_and_contact_from_fixed_content(self):
        # This repo's tests read real profile data via fixed_content_module()
        # rather than mocking it (see test_render_coverletter.py's
        # test_contact_info_comes_from_fixed_content) -- the active profile's
        # own contact info is asserted directly.
        render_coverletter_docx(_minimal_letter_data(), self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Alex Rivera", texts)
        contact_line = next(t for t in texts if "alex.rivera@example.com" in t)
        self.assertIn("PRODUCT MANAGER | GROWTH", contact_line)

    def test_recipient_block_uses_hiring_team_when_no_contact_name(self):
        render_coverletter_docx(_minimal_letter_data(), self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Acme Corp Hiring Team", texts)
        self.assertIn("Acme Corp", texts)

    def test_recipient_block_uses_attn_line_when_contact_name_present(self):
        render_coverletter_docx(
            _minimal_letter_data(
                contact_name="Maggie Smith", contact_title="HR Manager"
            ),
            self.out_path,
        )
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Attn: Maggie Smith, HR Manager", texts)

    def test_recipient_block_includes_location_when_present(self):
        render_coverletter_docx(
            _minimal_letter_data(company_location="Austin, TX"), self.out_path
        )
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Austin, TX", texts)

    def test_body_paragraphs_each_render_as_their_own_paragraph(self):
        render_coverletter_docx(_minimal_letter_data(), self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("I'm excited to apply for this role at Acme Corp.", texts)
        self.assertIn("My background lines up well with what you need.", texts)

    def test_sign_off_and_typed_name_render_with_no_embedded_image(self):
        render_coverletter_docx(_minimal_letter_data(), self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Sincerely,", texts)
        self.assertIn("Alex Rivera", texts)
        # Confirms the ATS-optimized "no signature image" design decision
        # (spec: docs/superpowers/specs/2026-08-17-docx-exporter-design.md) --
        # a real signature-image build would add an inline_shapes entry.
        self.assertEqual(len(doc.inline_shapes), 0)

    def test_returns_output_path(self):
        result = render_coverletter_docx(_minimal_letter_data(), self.out_path)
        self.assertEqual(result, self.out_path)

    def test_creates_parent_directory_if_missing(self):
        nested_path = os.path.join(
            os.path.dirname(__file__), "_tmp_docx_subdir", "letter.docx"
        )
        try:
            render_coverletter_docx(_minimal_letter_data(), nested_path)
            self.assertTrue(os.path.exists(nested_path))
        finally:
            if os.path.exists(nested_path):
                os.remove(nested_path)
            subdir = os.path.dirname(nested_path)
            if os.path.isdir(subdir):
                os.rmdir(subdir)

    def test_recipient_block_with_contact_name_without_title(self):
        render_coverletter_docx(
            _minimal_letter_data(contact_name="Maggie Smith", contact_title=""),
            self.out_path,
        )
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Attn: Maggie Smith", texts)

    def test_minimal_fields_without_tagline_greeting_or_signoff(self):
        data = {
            "company_name": "",
            "tagline": "",
            "greeting": "",
            "body_paragraphs": [],
            "sign_off": "",
        }
        render_coverletter_docx(data, self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Alex Rivera", texts)

    def test_main_cli_execution(self):
        import json
        import tempfile
        from unittest.mock import patch

        import render_coverletter_docx as rcl_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            json_path = os.path.join(tmpdir, "letter.json")
            out_docx = os.path.join(tmpdir, "letter.docx")
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(_minimal_letter_data(), f)
            with patch("sys.argv", ["render_coverletter_docx.py", json_path, out_docx]):
                rcl_docx.main()
                self.assertTrue(os.path.exists(out_docx))


if __name__ == "__main__":
    unittest.main()
