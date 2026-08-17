import os
import sys
import unittest

SCRIPTS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "scripts"
)
sys.path.insert(0, SCRIPTS_DIR)

from docx import Document  # noqa: E402
from render_resume_docx import render_resume_docx  # noqa: E402


def _minimal_resume_data(**overrides):
    data = {
        "NAME": "Jane Doe",
        "TAGLINE": "PRODUCT MANAGER | GROWTH",
        "PHONE": "555-123-4567",
        "EMAIL": "jane@example.com",
        "LINKEDIN_DISPLAY": "linkedin.com/in/janedoe",
        "LOCATION": "Austin, TX",
        "SUMMARY_TEXT": "<strong>Product leader with 10 years experience.</strong> Focused on growth.",
        "SKILLS": ["**Product:** Roadmapping, A/B Testing", "SQL"],
        "EXPERIENCE": [
            {
                "title": "Senior PM",
                "company": "Acme Corp",
                "location": "Remote",
                "period": "2020-Present",
                "achievements": ["Shipped a feature used by 1M users."],
            }
        ],
        "CERTIFICATIONS": [{"title": "PMP", "org": "PMI", "year": "2019"}],
        "EDUCATION": [
            {
                "degree": "B.S. Computer Science",
                "institution": "State University",
                "year": "2015",
            }
        ],
    }
    data.update(overrides)
    return data


class TestRenderResumeDocx(unittest.TestCase):

    def setUp(self):
        self.out_path = os.path.join(
            os.path.dirname(__file__), "_tmp_resume_docx_test.docx"
        )

    def tearDown(self):
        if os.path.exists(self.out_path):
            os.remove(self.out_path)

    def _paragraph_texts(self, doc):
        return [p.text for p in doc.paragraphs]

    def test_header_contains_name_and_contact_line(self):
        render_resume_docx(_minimal_resume_data(), self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Jane Doe", texts)
        contact_line = next(t for t in texts if "555-123-4567" in t)
        self.assertIn("PRODUCT MANAGER | GROWTH", contact_line)
        self.assertIn("jane@example.com", contact_line)
        self.assertIn("Austin, TX", contact_line)

    def test_summary_first_sentence_is_bold_and_strong_tags_are_stripped(self):
        render_resume_docx(_minimal_resume_data(), self.out_path)
        doc = Document(self.out_path)
        summary_para = next(p for p in doc.paragraphs if "Product leader" in p.text)
        self.assertNotIn("<strong>", summary_para.text)
        self.assertNotIn("</strong>", summary_para.text)
        self.assertTrue(summary_para.runs[0].bold)
        self.assertIn(
            "Product leader with 10 years experience.", summary_para.runs[0].text
        )
        self.assertFalse(summary_para.runs[-1].bold)

    def test_skills_markdown_bold_is_converted_to_a_bold_run(self):
        render_resume_docx(_minimal_resume_data(), self.out_path)
        doc = Document(self.out_path)
        skill_para = next(p for p in doc.paragraphs if "Roadmapping" in p.text)
        self.assertNotIn("**", skill_para.text)
        bold_run = next(r for r in skill_para.runs if r.text == "Product:")
        self.assertTrue(bold_run.bold)

    def test_experience_renders_title_meta_and_bulleted_achievements(self):
        render_resume_docx(_minimal_resume_data(), self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Senior PM", texts)
        meta_line = next(t for t in texts if "Acme Corp" in t)
        self.assertIn("Remote", meta_line)
        self.assertIn("2020-Present", meta_line)
        bullet_para = next(p for p in doc.paragraphs if "Shipped a feature" in p.text)
        self.assertEqual(bullet_para.style.name, "List Bullet")

    def test_size_revenue_is_appended_to_company_in_parentheses(self):
        data = _minimal_resume_data()
        data["EXPERIENCE"][0]["size_revenue"] = "500 employees"
        render_resume_docx(data, self.out_path)
        doc = Document(self.out_path)
        meta_line = next(p.text for p in doc.paragraphs if "Acme Corp" in p.text)
        self.assertIn("Acme Corp (500 employees)", meta_line)

    def test_clients_and_career_note_render_as_labeled_lines(self):
        data = _minimal_resume_data()
        data["EXPERIENCE"][0]["clients"] = "Fortune 500 retailers"
        data["EXPERIENCE"][0]["career_note"] = "Took a planned career break in 2019."
        render_resume_docx(data, self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        clients_line = next(t for t in texts if "Fortune 500 retailers" in t)
        self.assertIn("Clients:", clients_line)
        note_line = next(t for t in texts if "planned career break" in t)
        self.assertIn("Career Note:", note_line)

    def test_certifications_render_as_pipe_separated_line(self):
        render_resume_docx(_minimal_resume_data(), self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("PMP | PMI | 2019", texts)

    def test_education_renders_degree_meta_description_and_bullets(self):
        data = _minimal_resume_data()
        data["EDUCATION"][0]["description"] = "Focus in distributed systems."
        data["EDUCATION"][0]["bullets"] = ["Dean's List, all semesters."]
        render_resume_docx(data, self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        degree_line = next(t for t in texts if "B.S. Computer Science" in t)
        self.assertIn("State University", degree_line)
        self.assertIn("2015", degree_line)
        self.assertIn("Focus in distributed systems.", texts)
        bullet_para = next(p for p in doc.paragraphs if "Dean's List" in p.text)
        self.assertEqual(bullet_para.style.name, "List Bullet")

    def test_why_section_is_omitted_entirely_when_blank(self):
        render_resume_docx(
            _minimal_resume_data(WHY_TEXT="", SECTION_WHY=""), self.out_path
        )
        doc = Document(self.out_path)
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        self.assertNotIn("Additional Relevant Experience", headings)

    def test_why_section_is_omitted_when_literal_null_string(self):
        render_resume_docx(
            _minimal_resume_data(WHY_TEXT="null", SECTION_WHY="null"), self.out_path
        )
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertNotIn("null", [t.strip().lower() for t in texts])

    def test_why_section_renders_with_default_heading_and_strips_html_tags(self):
        data = _minimal_resume_data(
            WHY_TEXT="<p><em>I've long admired</em> this company's mission.</p><p>I'd love to contribute.</p>",
            SECTION_WHY="",
        )
        render_resume_docx(data, self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Additional Relevant Experience", texts)
        self.assertIn("I've long admired this company's mission.", texts)
        self.assertIn("I'd love to contribute.", texts)
        joined = "\n".join(texts)
        self.assertNotIn("<p>", joined)
        self.assertNotIn("<em>", joined)

    def test_why_section_uses_custom_heading_when_provided(self):
        data = _minimal_resume_data(
            WHY_TEXT="I'm a great fit.", SECTION_WHY="Why Acme Corp?"
        )
        render_resume_docx(data, self.out_path)
        doc = Document(self.out_path)
        texts = self._paragraph_texts(doc)
        self.assertIn("Why Acme Corp?", texts)

    def test_returns_output_path(self):
        result = render_resume_docx(_minimal_resume_data(), self.out_path)
        self.assertEqual(result, self.out_path)

    def test_creates_parent_directory_if_missing(self):
        nested_path = os.path.join(
            os.path.dirname(__file__), "_tmp_docx_subdir", "resume.docx"
        )
        try:
            render_resume_docx(_minimal_resume_data(), nested_path)
            self.assertTrue(os.path.exists(nested_path))
        finally:
            if os.path.exists(nested_path):
                os.remove(nested_path)
            subdir = os.path.dirname(nested_path)
            if os.path.isdir(subdir):
                os.rmdir(subdir)


if __name__ == "__main__":
    unittest.main()
