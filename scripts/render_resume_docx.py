"""
render_resume_docx.py — Builds an ATS-optimized .docx export of a tailored
resume directly from resume_data (the same uppercase-keyed dict
render_html.py consumes), using python-docx.

Usage (standalone):
    python scripts/render_resume_docx.py output/json/my_resume.json output/docx/my_resume.docx

Called programmatically by orchestrator.py's build_tailored_resume(), AFTER
the Step 7 trim-retry loop and its post-loop PDF text-layer check both
pass -- resume_data mutates across trim iterations (optional client rosters
dropped, the Why section dropped, LLM trim edits), so this must run on the
final, settled data, not the first successful PDF pass. See
docs/superpowers/specs/2026-08-17-docx-exporter-design.md.
"""

import argparse
import json
import os
import re

from docx import Document


def _add_bold_first_sentence(paragraph, text: str) -> None:
    """SUMMARY_TEXT wraps its first sentence in a literal <strong> tag
    (an HTML-rendering convention -- see render_html.py's own comment on
    SUMMARY_TEXT). Strip the tag and apply run-level bold instead, since
    docx has no inline-markup story."""
    match = re.match(r"<strong>(.*?)</strong>(.*)", text, re.DOTALL)
    if match:
        bold_part, rest = match.groups()
        run = paragraph.add_run(bold_part)
        run.bold = True
        if rest:
            paragraph.add_run(rest)
    else:
        paragraph.add_run(text)


def _add_bold_markdown_runs(paragraph, text: str) -> None:
    """SKILLS entries use "**Category:** Item, Item" markdown-style bold
    (the same convention render_html.py's build_skills_html() converts to
    <strong>) -- split on **...** and alternate bold/plain runs."""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = (i % 2 == 1)


def _is_blank_or_null(value: str) -> bool:
    return not value or value.strip().lower() == "null"


def render_resume_docx(resume_data: dict, output_path: str) -> str:
    """
    Builds an ATS-optimized .docx from resume_data and writes it to
    output_path. Returns output_path on success.
    """
    doc = Document()

    # --- Header ---
    doc.add_heading(resume_data.get("NAME", ""), level=0)
    contact_parts = [
        p for p in (
            resume_data.get("TAGLINE", ""),
            resume_data.get("PHONE", ""),
            resume_data.get("EMAIL", ""),
            resume_data.get("LINKEDIN_DISPLAY", ""),
            resume_data.get("LOCATION", ""),
        ) if p
    ]
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))

    # --- Summary ---
    summary_text = resume_data.get("SUMMARY_TEXT", "")
    if summary_text:
        doc.add_heading(resume_data.get("SECTION_SUMMARY", "Professional Summary"), level=1)
        p = doc.add_paragraph()
        _add_bold_first_sentence(p, summary_text)

    # --- Skills ---
    skills = resume_data.get("SKILLS", [])
    if skills:
        doc.add_heading(resume_data.get("SECTION_SKILLS", "Skills"), level=1)
        for skill in skills:
            p = doc.add_paragraph()
            _add_bold_markdown_runs(p, skill)

    # --- Experience ---
    experience = resume_data.get("EXPERIENCE", [])
    if experience:
        doc.add_heading(resume_data.get("SECTION_EXPERIENCE", "Work Experience"), level=1)
        for job in experience:
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(job.get("title", ""))
            title_run.bold = True

            company = job.get("company", "")
            if job.get("size_revenue"):
                company = f"{company} ({job['size_revenue']})"
            meta_parts = [p for p in (company, job.get("location", ""), job.get("period", "")) if p]
            if meta_parts:
                doc.add_paragraph(" | ".join(meta_parts))

            if job.get("clients"):
                p = doc.add_paragraph()
                run = p.add_run("Clients: ")
                run.bold = True
                p.add_run(job["clients"])

            for achievement in job.get("achievements", []):
                doc.add_paragraph(achievement, style="List Bullet")

            if job.get("career_note"):
                p = doc.add_paragraph()
                run = p.add_run("Career Note: ")
                run.bold = True
                p.add_run(job["career_note"])

    # --- Certifications ---
    certifications = resume_data.get("CERTIFICATIONS", [])
    if certifications:
        doc.add_heading(resume_data.get("SECTION_CERTIFICATIONS", "Training & Certifications"), level=1)
        for cert in certifications:
            cert_parts = [p for p in (cert.get("title", ""), cert.get("org", ""), cert.get("year", "")) if p]
            doc.add_paragraph(" | ".join(cert_parts))

    # --- Education ---
    education = resume_data.get("EDUCATION", [])
    if education:
        doc.add_heading(resume_data.get("SECTION_EDUCATION", "Education"), level=1)
        for edu in education:
            meta_parts = [m for m in (edu.get("institution", ""), edu.get("location", ""), edu.get("year", "")) if m]
            p = doc.add_paragraph()
            degree = edu.get("degree", "")
            if degree:
                run = p.add_run(degree)
                run.bold = True
            if meta_parts:
                meta_text = " | ".join(meta_parts)
                p.add_run(f" | {meta_text}" if degree else meta_text)
            if edu.get("description"):
                doc.add_paragraph(edu["description"])
            for bullet in edu.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # --- Why (optional) ---
    why_text = resume_data.get("WHY_TEXT", "")
    if not _is_blank_or_null(why_text):
        section_why = resume_data.get("SECTION_WHY", "")
        heading = section_why if not _is_blank_or_null(section_why) else "Additional Relevant Experience"
        doc.add_heading(heading, level=1)
        # WHY_TEXT contains literal <p>/<em> tags (see render_html.py's
        # build_why_html() comment) -- split into paragraphs on </p> and
        # strip the tags, rather than collapsing everything into one run.
        raw_paragraphs = [rp for rp in why_text.split("</p>") if rp.strip()]
        for raw_p in raw_paragraphs:
            clean = re.sub(r"</?p>|</?em>", "", raw_p).strip()
            if clean:
                doc.add_paragraph(clean)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Render a resume JSON file to .docx")
    parser.add_argument("input_json")
    parser.add_argument("output_docx")
    args = parser.parse_args()
    with open(args.input_json, "r", encoding="utf-8") as f:
        resume_data = json.load(f)
    render_resume_docx(resume_data, args.output_docx)
    print(f"Resume DOCX rendered -> {args.output_docx}")


if __name__ == "__main__":
    main()
