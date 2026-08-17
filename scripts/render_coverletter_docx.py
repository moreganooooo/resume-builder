"""
render_coverletter_docx.py — Builds an ATS-optimized .docx export of a
tailored cover letter directly from cover_letter_data (the same
lowercase-keyed dict render_coverletter.py consumes), using python-docx.

No embedded signature image (ATS-optimized fidelity call -- see
docs/superpowers/specs/2026-08-17-docx-exporter-design.md): typed name
only, matching what build_signature_block_html() degrades to anyway when a
profile has no signature.png.

Usage (standalone):
    python scripts/render_coverletter_docx.py output/json/my_letter_coverletter.json output/docx/my_letter_coverletter.docx

Called programmatically by orchestrator.py's build_tailored_coverletter().
"""

import argparse
import datetime
import json
import os

from docx import Document

import profile_paths


def _build_recipient_lines(company_name: str, contact_name: str = "", contact_title: str = "",
                            location: str = "") -> list[str]:
    """Same recipient-line logic as render_coverletter.py's
    build_recipient_block_html(), minus the HTML wrapping."""
    lines = []
    if contact_name:
        contact_line = f"Attn: {contact_name}"
        if contact_title:
            contact_line += f", {contact_title}"
        lines.append(contact_line)
    elif company_name:
        lines.append(f"{company_name} Hiring Team")
    if company_name:
        lines.append(company_name)
    if location:
        lines.append(location)
    return lines


def render_coverletter_docx(cover_letter_data: dict, output_path: str) -> str:
    """
    Builds an ATS-optimized .docx from cover_letter_data and writes it to
    output_path. Returns output_path on success.
    """
    contact = profile_paths.fixed_content_module().CONTACT_INFO
    doc = Document()

    # --- Header ---
    doc.add_heading(contact["NAME"], level=0)
    header_parts = [
        p for p in (
            cover_letter_data.get("tagline", ""),
            contact.get("PHONE", ""),
            contact.get("EMAIL", ""),
            contact.get("LINKEDIN_DISPLAY", ""),
            contact.get("LOCATION", ""),
        ) if p
    ]
    if header_parts:
        doc.add_paragraph(" | ".join(header_parts))

    # --- Date ---
    doc.add_paragraph(datetime.date.today().strftime("%B %-d, %Y"))

    # --- Recipient block ---
    for line in _build_recipient_lines(
        cover_letter_data.get("company_name", ""),
        cover_letter_data.get("contact_name", ""),
        cover_letter_data.get("contact_title", ""),
        cover_letter_data.get("company_location", ""),
    ):
        doc.add_paragraph(line)

    # --- Greeting ---
    greeting = cover_letter_data.get("greeting", "")
    if greeting:
        doc.add_paragraph(greeting)

    # --- Body ---
    for paragraph in cover_letter_data.get("body_paragraphs", []):
        doc.add_paragraph(paragraph)

    # --- Sign-off (typed name only -- no embedded signature image) ---
    sign_off = cover_letter_data.get("sign_off", "")
    if sign_off:
        doc.add_paragraph(sign_off)
    doc.add_paragraph(contact["NAME"])
    doc.add_paragraph(f"{contact['EMAIL']} | {contact['PHONE']}")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Render a cover letter JSON file to .docx")
    parser.add_argument("input_json")
    parser.add_argument("output_docx")
    args = parser.parse_args()
    with open(args.input_json, "r", encoding="utf-8") as f:
        cover_letter_data = json.load(f)
    render_coverletter_docx(cover_letter_data, args.output_docx)
    print(f"Cover letter DOCX rendered -> {args.output_docx}")


if __name__ == "__main__":
    main()
